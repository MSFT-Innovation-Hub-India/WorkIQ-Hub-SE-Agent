"""
Tool: create_word_doc
Create a Word document from agenda markdown content using python-docx.
"""

import logging
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

logger = logging.getLogger("hub_se_agent")

SCHEMA = {
    "type": "function",
    "name": "create_word_doc",
    "description": (
        "Create a Microsoft Word document from agenda markdown content. "
        "The document is saved to the configured output folder and opened automatically."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "filename": {
                "type": "string",
                "description": "The filename for the Word document (e.g. Agenda-Tesco-April-2026.docx).",
            },
            "markdown_content": {
                "type": "string",
                "description": "The full agenda markdown content including metadata header and table.",
            },
        },
        "required": ["filename", "markdown_content"],
    },
}


def _set_cell_borders(cell):
    """Add thin borders to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = tcBorders.makeelement(qn(f"w:{edge}"), {})
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")


def _set_run_font(run, size_pt, bold=False):
    """Set font properties on a run."""
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.bold = bold


def _add_formatted_text(cell, text, size_pt):
    """Add text to a cell, handling bullet points and line breaks."""
    cell.text = ""
    # Convert literal \n markers to actual newlines
    text = text.replace("\\n", "\n")
    # Split on line breaks within the cell
    lines = text.split("\n")
    first = True
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if first:
            para = cell.paragraphs[0]
            first = False
        else:
            para = cell.add_paragraph()
        # Check for bullet points
        stripped = line
        is_bullet = False
        if stripped.startswith("- "):
            stripped = stripped[2:]
            is_bullet = True
        elif stripped.startswith("* "):
            stripped = stripped[2:]
            is_bullet = True

        if is_bullet:
            para.style = "List Bullet"

        # Handle bold/italic markdown inline
        # Split on **bold** and *italic* markers
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", stripped)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                _set_run_font(run, size_pt, bold=True)
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                _set_run_font(run, size_pt)
                run.italic = True
            else:
                run = para.add_run(part)
                _set_run_font(run, size_pt)

        para.paragraph_format.space_after = Pt(2)


def _parse_markdown(markdown: str) -> tuple[dict, list[str], list[list[str]]]:
    """Parse agenda markdown into metadata, table headers, and table rows.

    Handles multi-line table rows where bullets/continuation lines don't
    start with '|'. Accumulates continuation lines into the last cell
    of the current row until a new row or non-table content is found.
    """
    metadata = {}
    headers = []
    rows = []

    lines = markdown.split("\n")
    in_table = False
    current_row: list[str] | None = None

    def _flush_row():
        nonlocal current_row
        if current_row is not None:
            rows.append(current_row)
            current_row = None

    # Pattern for a new table data row: starts with | followed by time
    _row_start = re.compile(r"^\|\s*\d{1,2}:\d{2}\s*(AM|PM)", re.IGNORECASE)
    # Pattern for structural rows (breaks, lunch, wrap-up) that start with | —
    _struct_row = re.compile(r"^\|\s*—\s*\|", re.IGNORECASE)

    for line in lines:
        line_stripped = line.strip()

        # Metadata: **Key:** Value
        m = re.match(r"\*\*(.+?):\*\*\s*(.*)", line_stripped)
        if m and not in_table:
            metadata[m.group(1).strip()] = m.group(2).strip()
            continue

        # Table header row
        if line_stripped.startswith("|") and "**Time**" in line_stripped:
            headers = [
                c.strip().strip("*").strip()
                for c in line_stripped.split("|")
                if c.strip()
            ]
            in_table = True
            continue

        # Skip separator row
        if in_table and re.match(r"^\|[-\s|]+\|$", line_stripped):
            continue

        # New data row (starts with | and a time like "09:00 AM")
        if in_table and (_row_start.match(line_stripped) or _struct_row.match(line_stripped)):
            _flush_row()
            cells = [c.strip() for c in line_stripped.split("|")]
            if cells and cells[0] == "":
                cells = cells[1:]
            if cells and cells[-1] == "":
                cells = cells[:-1]
            current_row = cells
            continue

        # Continuation line: part of the previous row's last cell
        if in_table and current_row is not None and line_stripped:
            # Strip trailing pipe if present (closing pipe of the row)
            cont = line_stripped.rstrip("|").strip()
            if cont:
                current_row[-1] = current_row[-1] + "\n" + cont
            continue

        # Non-table, non-empty line while in_table → end of table section
        # (but skip blank lines which are just spacing)
        if in_table and not line_stripped:
            continue

    _flush_row()
    return metadata, headers, rows


def handle(arguments: dict, *, on_progress=None, **kwargs) -> str:
    """Create Word document from markdown content."""
    filename = arguments["filename"]
    markdown_content = arguments["markdown_content"]

    # Get output folder and template from hub config
    try:
        import hub_config
        config = hub_config.load()
        output_folder = config.get("agenda_output_folder", "")
        template_path = config.get("agenda_template_path", "")
    except Exception:
        output_folder = ""
        template_path = ""

    if not output_folder:
        output_folder = str(Path.home() / "Documents" / "hub-se-agent-agenda-docs")

    output_path = Path(output_folder)
    output_path.mkdir(parents=True, exist_ok=True)
    file_path = output_path / filename

    if on_progress:
        on_progress("tool", f"Creating Word document: {filename}")

    try:
        metadata, headers, rows = _parse_markdown(markdown_content)

        # Open template if configured and exists; otherwise create blank doc
        use_template = False
        if template_path and Path(template_path).is_file():
            doc = Document(template_path)
            use_template = True
            logger.info("[WordDoc] Using template: %s", template_path)
        else:
            doc = Document()
            if template_path:
                logger.warning("[WordDoc] Template not found: %s — using blank document", template_path)

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        # If using template, add spacing after existing content (image/header)
        if use_template:
            doc.add_paragraph()

        # Title
        title_text = ""
        for line in markdown_content.split("\n"):
            if line.strip().startswith("# "):
                title_text = line.strip()[2:].strip()
                break
        if title_text:
            title = doc.add_heading(title_text, level=1)
            for run in title.runs:
                run.font.size = Pt(16)

        # Metadata block
        for key, value in metadata.items():
            para = doc.add_paragraph()
            run_key = para.add_run(f"{key}: ")
            _set_run_font(run_key, 12, bold=True)
            run_val = para.add_run(value)
            _set_run_font(run_val, 12)
            para.paragraph_format.space_after = Pt(2)

        # Day headings — find them in the markdown
        day_sections = re.split(r"(##\s+\*\*Day\s+\d+.*?\*\*)", markdown_content)

        # If there are day sections, process each; otherwise just do one table
        if len(day_sections) > 1:
            for i in range(1, len(day_sections), 2):
                day_heading = day_sections[i].strip().lstrip("#").strip().strip("*").strip()
                doc.add_paragraph()  # spacing
                h = doc.add_heading(day_heading, level=2)
                for run in h.runs:
                    run.font.size = Pt(13)

                day_content = day_sections[i + 1] if i + 1 < len(day_sections) else ""
                _, day_headers, day_rows = _parse_markdown(day_content)
                if not day_headers:
                    day_headers = headers
                _add_table(doc, day_headers, day_rows)
        else:
            if headers and rows:
                doc.add_paragraph()  # spacing
                _add_table(doc, headers, rows)

        doc.save(str(file_path))
        logger.info("[WordDoc] Created: %s", file_path)

        # Try to open the file
        if sys.platform == "win32":
            try:
                os.startfile(str(file_path))
            except Exception as e:
                logger.warning("[WordDoc] Could not auto-open: %s", e)

        return f"Document created successfully: {file_path}"

    except Exception as e:
        logger.error("[WordDoc] Failed to create document: %s", e, exc_info=True)
        return f"Error creating Word document: {e}"


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    """Add a formatted table to the document."""
    if not headers or not rows:
        return

    num_cols = len(headers)
    table = doc.add_table(rows=1, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Header row
    hdr = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(header_text)
        _set_run_font(run, 10, bold=True)
        # Grey background for header
        shading = cell._tc.get_or_add_tcPr().makeelement(
            qn("w:shd"),
            {qn("w:fill"): "D9E2F3", qn("w:val"): "clear"},
        )
        cell._tc.get_or_add_tcPr().append(shading)
        _set_cell_borders(cell)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i in range(num_cols):
            cell = row.cells[i]
            text = row_data[i] if i < len(row_data) else ""
            _add_formatted_text(cell, text, 10)
            _set_cell_borders(cell)

    # Set column widths — fix Time, Speaker, Topic; let Description fill remaining space
    if num_cols == 4:
        # Calculate available width from page width minus margins
        section = doc.sections[0]
        page_width = section.page_width - section.left_margin - section.right_margin
        fixed_widths = [Cm(3.5), Cm(3.5), Cm(4.5)]
        desc_width = page_width - sum(fixed_widths, Cm(0))
        widths = fixed_widths + [desc_width]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width
